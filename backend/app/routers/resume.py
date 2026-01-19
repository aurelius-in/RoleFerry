from fastapi import APIRouter, HTTPException, UploadFile, File
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import logging
import json
import io
import os
import re
import html as html_lib

from sqlalchemy import text as sql_text, bindparam
from sqlalchemy.dialects.postgresql import JSONB

from ..db import get_engine
from ..services_resume import parse_resume
from ..clients.openai_client import get_openai_client, extract_json_from_text
from ..storage import store


router = APIRouter()
engine = get_engine()
logger = logging.getLogger(__name__)
DEMO_USER_ID = "demo-user"
_ENABLE_DB_WRITES = os.getenv("ROLEFERRY_ENABLE_DB_WRITES", "false").lower() == "true"


def _looks_like_resume_has_metrics(text: str) -> bool:
    t = (text or "").lower()
    return bool(re.search(r"(\$|%|\b\d{1,3}\s*%|\b\d[\d,]*\s*(?:k|m|b)\+?\b|\b\d+\s*(?:million|billion)\b)", t, re.I))


def _looks_like_resume_has_experience(text: str) -> bool:
    t = (text or "").lower()
    return any(h in t for h in ["experience", "work history", "employment history", "professional experience"])


def _should_retry_resume_llm(raw_text: str, data: Dict[str, Any]) -> bool:
    """
    Decide whether to do a second LLM pass. We keep this conservative to avoid extra cost/latency.
    """
    if not isinstance(data, dict) or not data:
        return True
    pos = data.get("positions") or data.get("work_experience") or []
    skills = data.get("skills") or []
    kms = data.get("key_metrics") or data.get("keyMetrics") or []

    # If the resume clearly has experience and positions are missing, retry.
    if _looks_like_resume_has_experience(raw_text) and (not isinstance(pos, list) or len(pos) < 2):
        return True
    # If the resume has metrics and key_metrics is empty, retry.
    if _looks_like_resume_has_metrics(raw_text) and (not isinstance(kms, list) or len(kms) < 2):
        return True
    # If skills list is suspiciously tiny, retry.
    if not isinstance(skills, list) or len(skills) < 6:
        return True
    return False

class Position(BaseModel):
    company: str
    title: str
    start_date: str
    end_date: str
    current: bool
    description: str

class KeyMetric(BaseModel):
    metric: str
    value: str
    context: str

class Tenure(BaseModel):
    company: str
    duration: str
    role: str

class EducationItem(BaseModel):
    school: str
    degree: str
    field: str = ""
    start_year: str = ""
    end_year: str = ""
    notes: str = ""

class ResumeExtract(BaseModel):
    positions: List[Position]
    key_metrics: List[KeyMetric]
    business_challenges: List[str] = []
    skills: List[str]
    accomplishments: List[str]
    tenure: List[Tenure]
    education: List[EducationItem] = []


def _metric_from_line(s: str) -> KeyMetric:
    """
    Convert a best-effort metric line into {metric,value,context}.
    Keeps it deterministic and avoids inventing new facts.
    """
    t = str(s or "").strip().lstrip("-•* ").strip()
    if not t:
        return KeyMetric(metric="", value="", context="")

    # Find a likely "value" token first (%, $, k/m/b)
    m = re.search(r"(\$\s*\d[\d,]*(?:\.\d+)?\s*(?:[kKmMbB])?\+?|\b\d+%|\b\d[\d,]*(?:\.\d+)?\s*(?:k|m|b)\+?\b)", t, re.I)
    if not m:
        return KeyMetric(metric=t[:120], value="", context="")
    val = m.group(1).strip()
    before = t[: m.start()].strip(" :-—")
    after = t[m.end() :].strip(" :-—")
    metric = before[:80] if before else t[:80]
    # Common pattern: "... by 50%". Don't leave a dangling "by" in the metric field.
    metric = re.sub(r"\bby\s*$", "", metric.strip(), flags=re.I).strip()
    ctx = after[:140] if after else ""
    return KeyMetric(metric=metric[:120], value=val[:40], context=ctx[:160])

def _normalize_skill_tokens(skills: List[Any]) -> List[str]:
    """
    Clean skills from either deterministic or LLM output:
    - strip category labels like 'Languages:'
    - split embedded labels like 'Tools: Docker' -> Docker
    - split comma/pipe lists
    - drop obviously non-skill headings
    """
    if not isinstance(skills, list):
        return []
    out: List[str] = []
    seen = set()
    stop = {"languages", "frameworks", "libraries", "tools", "databases", "hosting", "operating systems", "skills"}

    def add(tok: str) -> None:
        t = (tok or "").strip().strip("•-* ").strip()
        if not t:
            return
        # Drop proficiency suffixes
        t = re.sub(r"\s*\([^)]*(?:professional|bilingual|fluent|beginner|intermediate|advanced)[^)]*\)\s*$", "", t, flags=re.I).strip()
        if not t:
            return
        low = t.lower()
        if low in stop or low.endswith(":"):
            return
        if len(t) > 60:
            return
        k = low
        if k in seen:
            return
        seen.add(k)
        out.append(t)

    for s in skills:
        if s is None:
            continue
        txt = str(s).strip()
        if not txt:
            continue
        # Split pipe/comma lists
        parts = []
        if "|" in txt:
            parts = [p.strip() for p in txt.split("|") if p.strip()]
        elif "," in txt and len(txt) <= 80:
            parts = [p.strip() for p in txt.split(",") if p.strip()]
        else:
            parts = [txt]
        for p in parts:
            # Strip leading category label
            m = re.match(r"^(?:languages?|frameworks?|libraries?|tools?|databases?|hosting|operating systems?|skills?)\b[^:]{0,40}:\s*(.+)$", p, flags=re.I)
            if m:
                add(m.group(1))
                continue
            # Embedded label: "X Tools: Docker" => Docker (and maybe keep X if it's a real skill)
            if ":" in p:
                rhs = p.split(":", 1)[1].strip()
                if rhs:
                    add(rhs)
                    continue
            add(p)
    return out[:80]


def _filter_llm_key_metrics(raw_text: str, key_metrics: List[KeyMetric]) -> List[KeyMetric]:
    """
    Keep only key metrics that are supported by explicit evidence in the resume text.
    This prevents the LLM from "helpfully" inventing numbers.
    """
    if not key_metrics:
        return []
    t = str(raw_text or "")
    t_low = t.lower()
    # Collect digit-like tokens found in the resume.
    numeric_tokens = set(re.findall(r"\d[\d,]*(?:\.\d+)?", t))
    out: List[KeyMetric] = []
    for km in key_metrics:
        metric = (km.metric or "").strip()
        value = (km.value or "").strip()
        context = (km.context or "").strip()
        if not (metric or value or context):
            continue
        # If value contains digits, require evidence in the resume text.
        vdigits = re.findall(r"\d[\d,]*(?:\.\d+)?", value)
        if vdigits:
            ok = False
            for v in vdigits:
                if v not in numeric_tokens:
                    continue
                # Require a standalone match for small integers to avoid accidental acceptance (e.g., years/dates).
                try:
                    v_int = int(v.replace(",", ""))
                except Exception:
                    v_int = None
                pattern = rf"\b{re.escape(v)}\b"
                m = re.search(pattern, t)
                if not m:
                    # If it's a large number (>= 1000), substring match is acceptable
                    if v_int is not None and v_int >= 1000 and v in t:
                        ok = True
                        break
                    continue
                # For smaller numbers, require a nearby keyword from metric/context.
                if v_int is not None and v_int < 1000:
                    window = t_low[max(0, m.start() - 60) : min(len(t_low), m.end() + 60)]
                    keywords = []
                    for w in re.findall(r"[a-z]{4,}", f"{metric} {context}".lower()):
                        if w not in {"with", "from", "that", "this", "into", "over", "used", "work", "role", "year", "years", "month", "months"}:
                            keywords.append(w)
                    if keywords and not any(k in window for k in keywords[:10]):
                        continue
                ok = True
                break
            if not ok:
                continue
        out.append(KeyMetric(metric=metric[:120], value=value[:40], context=context[:160]))
        if len(out) >= 12:
            break
    return out[:10]


def _filter_non_experience_positions(positions: List[Position]) -> List[Position]:
    """
    Remove items that look like education/training being misclassified as positions.
    """
    out: List[Position] = []
    for p in positions or []:
        title = (p.title or "").lower()
        company = (p.company or "").lower()
        if any(k in title for k in ["course", "curriculum", "immersive", "bootcamp", "boot camp", "certificate", "certification", "training"]):
            continue
        if any(k in company for k in ["university", "college", "school", "academy"]):
            continue
        out.append(p)
    return out


def _filter_non_experience_tenure(tenure: List[Tenure]) -> List[Tenure]:
    out: List[Tenure] = []
    for t in tenure or []:
        comp = (t.company or "").lower()
        role = (t.role or "").lower()
        if any(k in comp for k in ["university", "college", "school", "academy"]):
            continue
        if any(k in role for k in ["course", "curriculum", "bootcamp", "boot camp", "certificate", "certification", "training"]):
            continue
        out.append(t)
    return out
class ResumeExtractResponse(BaseModel):
    success: bool
    message: str
    extract: Optional[ResumeExtract] = None

@router.post("/upload", response_model=ResumeExtractResponse)
async def upload_resume(file: UploadFile = File(...)):
    """
    Upload and parse resume file to extract key information.
    """
    try:
        # Validate file type
        if not file.filename.lower().endswith((".pdf", ".docx", ".txt", ".html", ".htm")):
            raise HTTPException(status_code=400, detail="Only PDF, DOCX, TXT, and HTML files are supported")

        # Read file contents and extract text. We prefer true PDF/DOCX parsing;
        # if we can't extract meaningful text, we return an explicit error rather
        # than showing canned demo content (which confuses demos).
        contents = await file.read()

        def extract_text_from_pdf(data: bytes) -> str:
            try:
                from pypdf import PdfReader  # type: ignore
            except Exception as e:
                raise HTTPException(status_code=500, detail="PDF parsing dependency missing") from e
            reader = PdfReader(io.BytesIO(data))
            parts: List[str] = []
            for page in reader.pages:
                try:
                    t = page.extract_text() or ""
                except Exception:
                    t = ""
                if t.strip():
                    parts.append(t)
            return "\n\n".join(parts).strip()

        def extract_text_from_docx(data: bytes) -> str:
            try:
                import docx  # type: ignore
            except Exception as e:
                raise HTTPException(status_code=500, detail="DOCX parsing dependency missing") from e
            d = docx.Document(io.BytesIO(data))
            parts: List[str] = []
            for p in d.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())
            # Tables (some resumes are table-heavy)
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt = (cell.text or "").strip()
                        if txt:
                            parts.append(txt)
            return "\n".join(parts).strip()

        def extract_text_from_html(data: bytes) -> str:
            """
            Best-effort HTML -> text extraction (no external deps).
            Removes scripts/styles, converts common block tags to newlines, strips remaining tags,
            and decodes HTML entities.
            """
            try:
                s = data.decode("utf-8", errors="ignore")
            except Exception:
                s = ""
            if not s.strip():
                return ""
            # Remove script/style blocks
            s = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", s)
            s = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", s)
            # Replace common separators with newlines
            s = re.sub(r"(?i)<br\\s*/?>", "\n", s)
            s = re.sub(r"(?i)</p\\s*>", "\n", s)
            s = re.sub(r"(?i)</div\\s*>", "\n", s)
            s = re.sub(r"(?i)</li\\s*>", "\n", s)
            s = re.sub(r"(?i)</tr\\s*>", "\n", s)
            s = re.sub(r"(?i)</h[1-6]\\s*>", "\n", s)
            # Strip tags
            s = re.sub(r"(?is)<[^>]+>", " ", s)
            # Decode entities
            s = html_lib.unescape(s)
            # Normalize whitespace
            s = re.sub(r"[ \\t\\r\\f\\v]+", " ", s)
            s = re.sub(r"\\n\\s*\\n\\s*\\n+", "\n\n", s)
            return s.strip()

        filename = (file.filename or "").lower()
        raw_text = ""
        if filename.endswith(".pdf"):
            raw_text = extract_text_from_pdf(contents)
        elif filename.endswith(".docx"):
            raw_text = extract_text_from_docx(contents)
        elif filename.endswith((".html", ".htm")):
            raw_text = extract_text_from_html(contents)
        else:
            try:
                raw_text = contents.decode("utf-8", errors="ignore")
            except Exception:
                raw_text = ""

        # If we couldn't extract real text, stop and tell the user what happened.
        # (Common cause: scanned/image-only PDFs; requires OCR.)
        if not raw_text or len(raw_text.strip()) < 200:
            raise HTTPException(
                status_code=400,
                detail="Could not extract readable text from this file. If it's a scanned PDF, export as text-based PDF or upload a DOCX/TXT instead.",
            )

        # Run rule-based parser to populate resume.* style fields for storage
        parsed = parse_resume(raw_text or "")
        # Cache in memory so matching works even without DB
        store.demo_latest_resume = parsed
        store.demo_latest_resume_text = raw_text or ""

        # Persist to the RESUME table as raw text + parsed JSON for a demo user (best-effort).
        # For local dev it’s common to run without Postgres; DB writes are OFF by default.
        if _ENABLE_DB_WRITES:
            try:
                stmt = (
                    sql_text(
                        """
                        INSERT INTO resume (user_id, raw_text, parsed_json)
                        VALUES (:user_id, :raw_text, :parsed)
                        """
                    ).bindparams(bindparam("parsed", type_=JSONB))
                )
                async with engine.begin() as conn:
                    await conn.execute(
                        stmt,
                        {
                            "user_id": DEMO_USER_ID,
                            "raw_text": raw_text,
                            "parsed": parsed,
                        },
                    )
            except Exception:
                pass

        # Build a ResumeExtract for the UI. Prefer GPT-backed parsing when configured.
        # We no longer fall back to canned demo resume content here; if GPT fails,
        # we fall back to a minimal structure derived from the rule-based parser.
        client = get_openai_client()
        extract_obj: Optional[ResumeExtract] = None

        if client.should_use_real_llm and raw_text:
            try:
                raw = client.summarize_resume(raw_text)
                choices = raw.get("choices") or []
                msg = (choices[0].get("message") if choices else {}) or {}
                content_str = str(msg.get("content") or "")

                data = extract_json_from_text(content_str) or {}
                # If the model output is missing key fields, do one targeted retry.
                if _should_retry_resume_llm(raw_text, data):
                    try:
                        retry_messages = [
                            {
                                "role": "system",
                                "content": (
                                    "You are RoleFerry's resume parser. Your previous output was incomplete.\n"
                                    "Re-extract the resume into the REQUIRED JSON schema.\n\n"
                                    "Hard requirements:\n"
                                    "- Return ONLY a JSON object.\n"
                                    "- Do NOT invent facts.\n"
                                    "- Ensure positions covers the full work history with dates when present.\n"
                                    "- Ensure key_metrics includes measurable impacts if any numbers/percent/$ appear.\n"
                                    "- Ensure skills contains only real skills (no locations, degrees, awards/certs, headings, proficiency tags).\n\n"
                                    "Schema:\n"
                                    "{\n"
                                    '  \"positions\": [{\"company\":\"\",\"title\":\"\",\"start_date\":\"YYYY-MM|YYYY\",\"end_date\":\"YYYY-MM|YYYY\",\"current\":false,\"description\":\"\"}],\n'
                                    '  \"key_metrics\": [{\"metric\":\"\",\"value\":\"\",\"context\":\"\"}],\n'
                                    '  \"business_challenges\": [\"\"],\n'
                                    '  \"skills\": [\"\"],\n'
                                    '  \"accomplishments\": [\"\"],\n'
                                    '  \"tenure\": [{\"company\":\"\",\"duration\":\"\",\"role\":\"\"}]\n'
                                    "}\n"
                                ),
                            },
                            {"role": "user", "content": raw_text},
                        ]
                        raw2 = client.run_chat_completion(retry_messages, temperature=0.0, max_tokens=1600, stub_json={})
                        choices2 = raw2.get("choices") or []
                        msg2 = (choices2[0].get("message") if choices2 else {}) or {}
                        content_str2 = str(msg2.get("content") or "")
                        data2 = extract_json_from_text(content_str2) or {}
                        if isinstance(data2, dict) and data2:
                            data = data2
                    except Exception:
                        pass
                if isinstance(data, dict) and data:
                    positions_raw = data.get("positions") or data.get("work_experience") or []
                    tenure_raw = data.get("tenure") or []
                    key_metrics_raw = data.get("key_metrics") or data.get("keyMetrics") or []
                    accomplishments_raw = data.get("accomplishments") or []
                    skills_raw = data.get("skills") or []
                    business_challenges_raw = (
                        data.get("business_challenges")
                        or data.get("businessChallenges")
                        or data.get("business_challenges_solved")
                        or data.get("problems_solved")
                        or []
                    )

                    positions: List[Position] = []
                    for p in positions_raw:
                        if not isinstance(p, dict):
                            continue
                        positions.append(
                            Position(
                                company=str(p.get("company") or ""),
                                title=str(p.get("title") or ""),
                                start_date=str(p.get("start_date") or ""),
                                end_date=str(p.get("end_date") or ""),
                                current=bool(p.get("current") or False),
                                description=str(p.get("description") or ""),
                            )
                        )

                    tenure: List[Tenure] = []
                    for t in tenure_raw:
                        if not isinstance(t, dict):
                            continue
                        tenure.append(
                            Tenure(
                                company=str(t.get("company") or ""),
                                duration=str(t.get("duration") or ""),
                                role=str(t.get("role") or ""),
                            )
                        )

                    key_metrics: List[KeyMetric] = []
                    for m in key_metrics_raw:
                        if isinstance(m, dict):
                            key_metrics.append(
                                KeyMetric(
                                    metric=str(m.get("metric") or ""),
                                    value=str(m.get("value") or ""),
                                    context=str(m.get("context") or ""),
                                )
                            )
                        else:
                            key_metrics.append(
                                KeyMetric(metric=str(m), value="", context="")
                            )
                    key_metrics = _filter_llm_key_metrics(raw_text, key_metrics)

                    accomplishments = [str(a) for a in accomplishments_raw]
                    skills = _normalize_skill_tokens([str(s) for s in skills_raw])
                    business_challenges = [str(b) for b in business_challenges_raw]
                    education_raw = data.get("education") or []
                    education: List[EducationItem] = []
                    for e in education_raw if isinstance(education_raw, list) else []:
                        if isinstance(e, dict):
                            education.append(
                                EducationItem(
                                    school=str(e.get("school") or ""),
                                    degree=str(e.get("degree") or ""),
                                    field=str(e.get("field") or ""),
                                    start_year=str(e.get("start_year") or ""),
                                    end_year=str(e.get("end_year") or ""),
                                    notes=str(e.get("notes") or ""),
                                )
                            )

                    extract_obj = ResumeExtract(
                        positions=positions,
                        key_metrics=key_metrics,
                        business_challenges=business_challenges,
                        skills=skills,
                        accomplishments=accomplishments,
                        tenure=tenure,
                        education=education,
                    )
            except Exception:
                # On any GPT failure, we'll fall back to the rule-based parser output below.
                extract_obj = None

        # If GPT succeeded but missed fields, backfill from rule-based parsing
        if extract_obj is not None:
            # If GPT produced weak/garbled positions (common on some DOCX formats),
            # replace with the rule-based positions/tenure derived from the same resume text.
            try:
                bad_positions = False
                if not extract_obj.positions:
                    bad_positions = True
                else:
                    # Heuristic: if most titles/companies are empty or look like long sentences, treat as bad.
                    bad = 0
                    for p in extract_obj.positions[:8]:
                        title = (p.title or "").strip()
                        company = (p.company or "").strip()
                        if not title or not company:
                            bad += 1
                            continue
                        if len(title.split()) > 12:
                            bad += 1
                    if bad >= max(2, len(extract_obj.positions[:8]) // 2):
                        bad_positions = True

                if bad_positions:
                    pos_raw = parsed.get("Positions") or []
                    positions: List[Position] = []
                    for p in pos_raw[:10] if isinstance(pos_raw, list) else []:
                        if isinstance(p, dict):
                            positions.append(
                                Position(
                                    company=str(p.get("company") or p.get("Company") or ""),
                                    title=str(p.get("title") or p.get("Title") or ""),
                                    start_date=str(p.get("start_date") or p.get("StartDate") or ""),
                                    end_date=str(p.get("end_date") or p.get("EndDate") or ""),
                                    current=bool(p.get("current") or False),
                                    description=str(p.get("description") or p.get("Description") or ""),
                                )
                            )
                    if positions:
                        extract_obj.positions = _filter_non_experience_positions(positions)[:8]

                    tenure_raw = parsed.get("Tenure") or []
                    if isinstance(tenure_raw, list) and tenure_raw:
                        extract_obj.tenure = [
                            Tenure(
                                company=str(t.get("company") or ""),
                                duration=str(t.get("duration") or ""),
                                role=str(t.get("role") or ""),
                            )
                            for t in tenure_raw
                            if isinstance(t, dict)
                        ][:12]
            except Exception:
                pass

            if not extract_obj.key_metrics:
                km_raw = parsed.get("KeyMetrics") or parsed.get("key_metrics") or []
                key_metrics: List[KeyMetric] = []
                for m in km_raw[:12] if isinstance(km_raw, list) else []:
                    if isinstance(m, dict):
                        key_metrics.append(
                            KeyMetric(
                                metric=str(m.get("metric") or m.get("Metric") or ""),
                                value=str(m.get("value") or m.get("Value") or ""),
                                context=str(m.get("context") or m.get("Context") or ""),
                            )
                        )
                    else:
                        km = _metric_from_line(str(m))
                        if km.metric:
                            key_metrics.append(km)
                extract_obj.key_metrics = key_metrics[:10]
            else:
                # Even if LLM provided metrics, filter out unsupported/invented numbers.
                filtered = _filter_llm_key_metrics(raw_text, list(extract_obj.key_metrics or []))
                if filtered:
                    extract_obj.key_metrics = filtered[:10]
                else:
                    # Prefer deterministic if LLM metrics were not supported by evidence.
                    km_raw = parsed.get("KeyMetrics") or parsed.get("key_metrics") or []
                    key_metrics: List[KeyMetric] = []
                    for m in km_raw[:12] if isinstance(km_raw, list) else []:
                        if isinstance(m, dict):
                            key_metrics.append(
                                KeyMetric(
                                    metric=str(m.get("metric") or m.get("Metric") or ""),
                                    value=str(m.get("value") or m.get("Value") or ""),
                                    context=str(m.get("context") or m.get("Context") or ""),
                                )
                            )
                        else:
                            km = _metric_from_line(str(m))
                            if km.metric:
                                key_metrics.append(km)
                    extract_obj.key_metrics = key_metrics[:10]
            if not extract_obj.skills:
                extract_obj.skills = _normalize_skill_tokens([str(s) for s in (parsed.get("Skills") or [])])
            if not extract_obj.business_challenges:
                extract_obj.business_challenges = [
                    str(x) for x in (parsed.get("BusinessChallengesSolved") or parsed.get("ProblemsSolved") or [])
                ]
            if not extract_obj.accomplishments:
                extract_obj.accomplishments = [str(a) for a in (parsed.get("NotableAccomplishments") or [])]
            if not extract_obj.tenure:
                tenure_raw = parsed.get("Tenure") or []
                if isinstance(tenure_raw, list):
                    extract_obj.tenure = [
                        Tenure(
                            company=str(t.get("company") or ""),
                            duration=str(t.get("duration") or ""),
                            role=str(t.get("role") or ""),
                        )
                        for t in tenure_raw
                        if isinstance(t, dict)
                    ]
            if not extract_obj.education:
                edu_raw = parsed.get("Education") or []
                if isinstance(edu_raw, list):
                    extract_obj.education = [
                        EducationItem(
                            school=str(e.get("school") or ""),
                            degree=str(e.get("degree") or ""),
                            field=str(e.get("field") or ""),
                            start_year=str(e.get("start_year") or ""),
                            end_year=str(e.get("end_year") or ""),
                            notes=str(e.get("notes") or ""),
                        )
                        for e in edu_raw
                        if isinstance(e, dict)
                    ][:10]
            # If we still couldn't compute tenure, derive a minimal tenure list from positions.
            if not extract_obj.tenure and extract_obj.positions:
                extract_obj.tenure = [
                    Tenure(company=p.company, duration="", role=p.title)
                    for p in extract_obj.positions[:10]
                    if (p.company or p.title)
                ]

            # Always filter out education/training misclassified as experience (from either GPT or deterministic backfill).
            try:
                extract_obj.positions = _filter_non_experience_positions(list(extract_obj.positions or []))
                extract_obj.tenure = _filter_non_experience_tenure(list(extract_obj.tenure or []))
            except Exception:
                pass

        if extract_obj is None:
            # Rule-based fallback derived from the parsed resume (best-effort).
            # This should still reflect the uploaded resume text, not canned demo data.
            positions: List[Position] = []
            # NOTE: parse_resume() returns "Positions" (capital-P). Keep other keys for backwards compat.
            for p in (parsed.get("WorkExperience") or parsed.get("Positions") or parsed.get("positions") or [])[:6]:
                if isinstance(p, dict):
                    positions.append(
                        Position(
                            company=str(p.get("company") or p.get("Company") or ""),
                            title=str(p.get("title") or p.get("Title") or ""),
                            start_date=str(p.get("start_date") or p.get("StartDate") or ""),
                            end_date=str(p.get("end_date") or p.get("EndDate") or ""),
                            current=bool(p.get("current") or False),
                            description=str(p.get("description") or p.get("Description") or ""),
                        )
                    )

            key_metrics: List[KeyMetric] = []
            for m in (parsed.get("KeyMetrics") or parsed.get("key_metrics") or [])[:10]:
                if isinstance(m, dict):
                    key_metrics.append(
                        KeyMetric(
                            metric=str(m.get("metric") or m.get("Metric") or ""),
                            value=str(m.get("value") or m.get("Value") or ""),
                            context=str(m.get("context") or m.get("Context") or ""),
                        )
                    )
                else:
                    km = _metric_from_line(str(m))
                    if km.metric:
                        key_metrics.append(km)

            skills = _normalize_skill_tokens([str(s) for s in (parsed.get("Skills") or parsed.get("skills") or [])])
            business_challenges = [
                str(x) for x in (parsed.get("BusinessChallengesSolved") or parsed.get("ProblemsSolved") or [])
            ][:15]
            accomplishments = [str(a) for a in (parsed.get("NotableAccomplishments") or parsed.get("accomplishments") or [])][:25]
            tenure: List[Tenure] = []
            for t in (parsed.get("Tenure") or parsed.get("tenure") or [])[:10]:
                if isinstance(t, dict):
                    tenure.append(
                        Tenure(
                            company=str(t.get("company") or t.get("Company") or ""),
                            duration=str(t.get("duration") or t.get("Duration") or ""),
                            role=str(t.get("role") or t.get("Role") or ""),
                        )
                    )

            education: List[EducationItem] = []
            for e in (parsed.get("Education") or [])[:10]:
                if isinstance(e, dict):
                    education.append(
                        EducationItem(
                            school=str(e.get("school") or ""),
                            degree=str(e.get("degree") or ""),
                            field=str(e.get("field") or ""),
                            start_year=str(e.get("start_year") or ""),
                            end_year=str(e.get("end_year") or ""),
                            notes=str(e.get("notes") or ""),
                        )
                    )

            extract_obj = ResumeExtract(
                positions=positions,
                key_metrics=key_metrics,
                business_challenges=business_challenges,
                skills=skills,
                accomplishments=accomplishments,
                tenure=tenure,
                education=education,
            )
            # Filter out education/training misclassified as experience.
            try:
                extract_obj.positions = _filter_non_experience_positions(list(extract_obj.positions or []))
                extract_obj.tenure = _filter_non_experience_tenure(list(extract_obj.tenure or []))
            except Exception:
                pass

        return ResumeExtractResponse(
            success=True,
            message="Resume parsed successfully",
            extract=extract_obj,
        )
    except HTTPException:
        # Preserve specific status codes/messages (e.g., unreadable/scanned PDFs)
        raise
    except Exception:
        logger.exception("Error parsing resume")
        raise HTTPException(status_code=500, detail="Failed to parse resume")

@router.post("/save", response_model=ResumeExtractResponse)
async def save_resume_extract(extract: ResumeExtract):
    """
    Save resume extract for a user.
    """
    try:
        # In a real app, save to database with user_id
        return ResumeExtractResponse(
            success=True,
            message="Resume extract saved successfully",
            extract=extract
        )
    except Exception as e:
        logger.exception("Error saving resume extract")
        raise HTTPException(status_code=500, detail="Failed to save resume extract")

@router.get("/{user_id}", response_model=ResumeExtractResponse)
async def get_resume_extract(user_id: str):
    """
    Get resume extract for a user.
    """
    try:
        # In a real app, fetch from database
        # For now, return mock data
        mock_extract = ResumeExtract(
            positions=[
                Position(
                    company="TechCorp Inc.",
                    title="Senior Software Engineer",
                    start_date="2022-01",
                    end_date="2024-12",
                    current=True,
                    description="Led development of microservices architecture, reducing system latency by 40%"
                )
            ],
            key_metrics=[
                KeyMetric(
                    metric="System Performance",
                    value="40% reduction",
                    context="in latency through microservices optimization"
                )
            ],
            skills=["Python", "JavaScript", "React", "Node.js", "AWS", "Docker", "PostgreSQL"],
            accomplishments=[
                "Reduced system latency by 40% through microservices architecture"
            ],
            tenure=[
                Tenure(company="TechCorp Inc.", duration="2 years", role="Senior Software Engineer")
            ]
        )
        
        return ResumeExtractResponse(
            success=True,
            message="Resume extract retrieved successfully",
            extract=mock_extract
        )
    except Exception as e:
        logger.exception("Error retrieving resume extract")
        raise HTTPException(status_code=500, detail="Failed to get resume extract")

@router.put("/{user_id}", response_model=ResumeExtractResponse)
async def update_resume_extract(user_id: str, extract: ResumeExtract):
    """
    Update resume extract for a user.
    """
    try:
        # In a real app, update in database
        return ResumeExtractResponse(
            success=True,
            message="Resume extract updated successfully",
            extract=extract
        )
    except Exception as e:
        logger.exception("Error updating resume extract")
        raise HTTPException(status_code=500, detail="Failed to update resume extract")

@router.delete("/{user_id}")
async def delete_resume_extract(user_id: str):
    """
    Delete resume extract for a user.
    """
    try:
        # In a real app, delete from database
        return {"success": True, "message": "Resume extract deleted successfully"}
    except Exception as e:
        logger.exception("Error deleting resume extract")
        raise HTTPException(status_code=500, detail="Failed to delete resume extract")
